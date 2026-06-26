#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
会议文档生成器 - 独立可迁移模块

功能：
  1. 笔记生成：HTML / Markdown / Editor.js JSON → DOCX
  2. 纪要生成：支持模板（默认模板/会议纪要模板/联合行文模板）+ 白板模式
  3. 压缩包打包：多文件 → ZIP

依赖：
  pip install python-docx beautifulsoup4 markdown loguru

用法示例：
  from meeting_doc_gen import MeetingDocGenerator

  gen = MeetingDocGenerator(template_dir="./templates", output_dir="./output")

  # 生成笔记
  success, result, err = gen.build_note(
      content="<p>会议讨论了项目进度...</p>",
      title="项目周会 - 笔记",
      content_type="html"
  )

  # 生成纪要（使用模板）
  success, result, err = gen.build_summary(
      content="<h2>会议纪要</h2><p>讨论内容...</p>",
      meeting_data={
          "theme": "项目周会",
          "recording_time": 1719123456000,
          "places": "3楼会议室",
          "moderator": "张三",
          "attendees": ["李四", "王五"],
          "kind": 0,
          "remark": ""
      },
      template_type=2,
      content_type="html"
  )

  # 打包为 ZIP
  success, zip_path, err = gen.package_to_zip(
      files=[("笔记.docx", b"..."), ("纪要.docx", b"...")],
      zip_name="项目周会_文档包"
  )
"""

import json
import os
import re
import zipfile
from datetime import datetime
from typing import Any, Dict, List, Optional, Tuple, Union

from docx import Document
from loguru import logger

from .models import Alignment, DocumentNode, ParagraphNode
from .config import (
    KIND_LABEL_MAP, MEETING_MAPPING, apply_mapping, format_value, _fmt_datetime,
)
from .html_parser import is_html_content, is_editorjs_format
from .doc_builder import build_from_html, build_from_markdown
from .docx_renderer import render_document, save_docx
from .editorjs_parser import EditorJsParser
from .template_engine import load_template, replace_placeholders, replace_rich_placeholder


# =============================================================================
# Editor.js 便捷函数
# =============================================================================

def build_from_editor_js(editorjs_data: Union[str, Dict],
                         output_path: Optional[str] = None,
                         title: Optional[str] = None) -> Tuple[bool, Any, str]:
    """将 Editor.js JSON 转换为 DOCX 文档（不使用模板）"""
    try:
        parser = EditorJsParser()
        doc_node = parser.parse(editorjs_data)

        if title:
            section = doc_node.sections[0] if doc_node.sections else doc_node.add_section()
            title_para = ParagraphNode(alignment=Alignment.CENTER, line_spacing=1.0, space_after=16)
            title_para.add_run(title, font_name='方正小标宋简体', font_size=22, bold=False)
            section.body.insert(0, title_para)

        doc = Document()
        render_document(doc, doc_node)
        return save_docx(doc, output_path)
    except Exception as e:
        logger.error(f"Editor.js 转换失败: {e}")
        return False, None, str(e)


def convert_editorjs_with_template(editorjs_data: Union[str, Dict],
                                   template_path: str,
                                   placeholders: Optional[Dict[str, str]] = None,
                                   rich_placeholders: Optional[Dict] = None,
                                   output_path: Optional[str] = None) -> Tuple[bool, Any, str]:
    """使用模板将 Editor.js JSON 转换为 DOCX"""
    try:
        doc = load_template(template_path)
        parser = EditorJsParser()
        doc_node = parser.parse(editorjs_data)

        if placeholders:
            replace_placeholders(doc, placeholders)

        if rich_placeholders:
            for placeholder, content in rich_placeholders.items():
                if isinstance(content, str):
                    try:
                        content_data = json.loads(content) if content.strip().startswith('{') else None
                        if content_data:
                            content_node = parser.parse(content_data)
                            nodes = []
                            for section in content_node.sections:
                                nodes.extend(section.body)
                            replace_rich_placeholder(doc, placeholder, nodes)
                        else:
                            para = ParagraphNode(line_spacing=1.0)
                            para.add_run(content, font_name='仿宋', font_size=16)
                            replace_rich_placeholder(doc, placeholder, [para])
                    except Exception:
                        para = ParagraphNode(line_spacing=1.0)
                        para.add_run(content, font_name='仿宋', font_size=16)
                        replace_rich_placeholder(doc, placeholder, [para])
                elif isinstance(content, dict):
                    content_node = parser.parse(content)
                    nodes = []
                    for section in content_node.sections:
                        nodes.extend(section.body)
                    replace_rich_placeholder(doc, placeholder, nodes)
                elif isinstance(content, list):
                    replace_rich_placeholder(doc, placeholder, content)

        if not rich_placeholders:
            for section in doc_node.sections:
                for node in section.body:
                    if isinstance(node, ParagraphNode):
                        new_para = doc.add_paragraph()
                        from .docx_renderer import apply_paragraph_format, render_run
                        apply_paragraph_format(new_para, node)
                        for run_node in node.runs:
                            render_run(new_para, run_node)
                    else:
                        from .docx_renderer import render_table
                        render_table(doc, node)

        return save_docx(doc, output_path)
    except Exception as e:
        logger.error(f"模板转换失败: {e}")
        return False, None, str(e)


# =============================================================================
# 主类：MeetingDocGenerator
# =============================================================================

class MeetingDocGenerator:
    """
    会议文档生成器

    提供笔记生成、纪要生成、ZIP 打包三大核心功能。

    Attributes:
        template_dir: 模板文件目录路径
        output_dir: 输出文件目录路径
    """

    TEMPLATE_MAP = {1: '默认模板', 2: '会议纪要模板', 3: '联合行文模板'}

    def __init__(self, template_dir: str = "./templates", output_dir: str = "./output"):
        """
        初始化生成器

        Args:
            template_dir: 模板文件目录，结构如下：
                templates/
                ├── 默认模板.docx
                ├── note/
                │   └── 默认笔记模板.docx
                └── summary/
                    ├── 会议纪要模板.docx
                    └── 联合行文模板.docx
            output_dir: 输出目录
        """
        self.template_dir = template_dir
        self.output_dir = output_dir
        os.makedirs(output_dir, exist_ok=True)
        self._templates = self._scan_templates()

    def _scan_templates(self) -> Dict[str, Dict[str, str]]:
        """扫描模板目录"""
        templates = {'_root': {}, 'note': {}, 'summary': {}}
        if not os.path.isdir(self.template_dir):
            return templates
        try:
            for f in os.listdir(self.template_dir):
                fpath = os.path.join(self.template_dir, f)
                if os.path.isfile(fpath) and f.lower().endswith('.docx') and not f.startswith('~'):
                    templates['_root'][os.path.splitext(f)[0]] = fpath
            for category in ['note', 'summary']:
                cat_dir = os.path.join(self.template_dir, category)
                if os.path.isdir(cat_dir):
                    for f in os.listdir(cat_dir):
                        fpath = os.path.join(cat_dir, f)
                        if os.path.isfile(fpath) and f.lower().endswith('.docx') and not f.startswith('~'):
                            templates[category][os.path.splitext(f)[0]] = fpath
        except Exception as e:
            logger.warning(f"扫描模板目录失败: {e}")
        return templates

    def get_template_path(self, name: str, category: Optional[str] = None) -> Optional[str]:
        """获取模板文件路径"""
        return self._templates.get(category or '_root', {}).get(name)

    # -------------------------------------------------------------------------
    # 笔记生成
    # -------------------------------------------------------------------------

    def build_note(self, content: str, title: Optional[str] = None,
                   content_type: str = "auto",
                   output_path: Optional[str] = None) -> Tuple[bool, Any, str]:
        """
        生成笔记文档

        Args:
            content: 笔记内容（HTML / Markdown / Editor.js JSON / 纯文本）
            title: 文档标题
            content_type: "auto" / "html" / "markdown" / "editorjs" / "text"
            output_path: 输出路径（不含扩展名），为 None 时返回 bytes

        Returns:
            (success, result, error)
        """
        if not content or not content.strip():
            return False, None, "笔记内容为空"

        try:
            content_type = self._detect_content_type(content, content_type)

            if content_type == "editorjs":
                return build_from_editor_js(content, output_path, title)
            elif content_type == "markdown":
                doc_node = build_from_markdown(content, title)
            elif content_type == "text":
                paragraphs = [f"<p>{line}</p>" for line in content.split('\n') if line.strip()]
                doc_node = build_from_html(''.join(paragraphs), title)
            else:
                doc_node = build_from_html(content, title)

            doc = Document()
            render_document(doc, doc_node)
            return save_docx(doc, output_path)

        except Exception as e:
            logger.error(f"生成笔记失败: {e}")
            return False, None, str(e)

    # -------------------------------------------------------------------------
    # 纪要生成
    # -------------------------------------------------------------------------

    def build_summary(self, content: str,
                      meeting_data: Optional[Dict[str, Any]] = None,
                      template_type: int = 1,
                      content_type: str = "auto",
                      output_path: Optional[str] = None) -> Tuple[bool, Any, str]:
        """
        生成会议纪要文档

        Args:
            content: 纪要正文内容
            meeting_data: 会议元数据，字段：
                theme, recording_time(ms), places, moderator,
                attendees(list), kind(0-3), remark
            template_type: 1=白板 2=会议纪要模板 3=联合行文模板
            content_type: "auto" / "html" / "markdown" / "editorjs"
            output_path: 输出路径

        Returns:
            (success, result, error)
        """
        if not content or not content.strip():
            return False, None, "纪要内容为空"

        meeting_data = meeting_data or {}

        try:
            content_type = self._detect_content_type(content, content_type)
            content = self._preprocess_summary_content(content, content_type)

            # 准备会议信息
            attendees_list = meeting_data.get('attendees', [])
            if isinstance(attendees_list, str):
                try:
                    attendees_list = json.loads(attendees_list)
                except json.JSONDecodeError:
                    attendees_list = [attendees_list] if attendees_list else []

            recording_time_str = ''
            rt = meeting_data.get('recording_time')
            if rt:
                recording_time_str = _fmt_datetime(rt)

            attendees_str = '、'.join(attendees_list) if attendees_list else '无'

            # 白板模式
            if template_type == 1:
                return self._build_summary_plain(content, meeting_data, content_type,
                                                 recording_time_str, attendees_str, output_path)

            # 模板模式
            template_name = self.TEMPLATE_MAP.get(template_type)
            if not template_name:
                return False, None, f"不支持的模板类型: {template_type}"

            template_path = self.get_template_path(template_name, 'summary')
            if not template_path or not os.path.exists(template_path):
                return False, None, f"模板文件不存在: {template_name}"

            if content_type == "editorjs":
                return self._build_summary_editorjs_template(
                    content, template_path, meeting_data, recording_time_str, attendees_str, output_path)
            else:
                return self._build_summary_html_template(
                    content, template_path, meeting_data, recording_time_str, attendees_str, output_path)

        except Exception as e:
            logger.error(f"生成纪要失败: {e}")
            return False, None, str(e)

    # -------------------------------------------------------------------------
    # ZIP 打包
    # -------------------------------------------------------------------------

    def package_to_zip(self, files: List[Tuple[str, bytes]],
                       zip_name: str = "文档包",
                       output_dir: Optional[str] = None) -> Tuple[bool, str, str]:
        """
        将多个文件打包为 ZIP

        Args:
            files: [(文件名, bytes), ...]
            zip_name: ZIP 文件名（不含扩展名）
            output_dir: 输出目录

        Returns:
            (success, zip_path, error)
        """
        if not files:
            return False, '', "没有可打包的文件"

        output_dir = output_dir or self.output_dir
        os.makedirs(output_dir, exist_ok=True)

        safe_name = re.sub(r'[\x00-\x1F\x7F:/\\*?"<>|]', '_', zip_name)
        zip_path = os.path.join(output_dir, f"{safe_name}.zip")

        try:
            with zipfile.ZipFile(zip_path, 'w', zipfile.ZIP_DEFLATED, compresslevel=1) as zf:
                for file_name, file_content in files:
                    zf.writestr(file_name, file_content)
            return True, zip_path, ''
        except Exception as e:
            logger.error(f"打包 ZIP 失败: {e}")
            return False, '', str(e)

    def build_and_package(self, note_content: Optional[str] = None,
                          summary_content: Optional[str] = None,
                          meeting_data: Optional[Dict[str, Any]] = None,
                          template_type: int = 1,
                          content_type: str = "auto",
                          zip_name: str = "会议文档包",
                          output_dir: Optional[str] = None) -> Tuple[bool, str, str]:
        """
        一站式生成笔记+纪要并打包为 ZIP

        Args:
            note_content: 笔记内容，None 则不生成
            summary_content: 纪要内容，None 则不生成
            meeting_data: 会议元数据
            template_type: 纪要模板类型
            content_type: 内容类型
            zip_name: ZIP 文件名
            output_dir: 输出目录

        Returns:
            (success, zip_path, error)
        """
        files_to_pack = []

        if note_content:
            success, result, err = self.build_note(note_content, title="笔记", content_type=content_type)
            if not success:
                return False, '', f"生成笔记失败: {err}"
            note_bytes = result if isinstance(result, bytes) else open(result, 'rb').read()
            files_to_pack.append(("笔记.docx", note_bytes))

        if summary_content:
            success, result, err = self.build_summary(
                summary_content, meeting_data=meeting_data,
                template_type=template_type, content_type=content_type
            )
            if not success:
                return False, '', f"生成纪要失败: {err}"
            summary_bytes = result if isinstance(result, bytes) else open(result, 'rb').read()
            files_to_pack.append(("会议纪要.docx", summary_bytes))

        if not files_to_pack:
            return False, '', "没有可打包的内容"

        return self.package_to_zip(files_to_pack, zip_name, output_dir)

    # -------------------------------------------------------------------------
    # 内部辅助方法
    # -------------------------------------------------------------------------

    def _detect_content_type(self, content: str, content_type: str) -> str:
        """自动检测内容类型"""
        if content_type != "auto":
            return content_type
        if is_editorjs_format(content):
            return "editorjs"
        if is_html_content(content):
            return "html"
        return "text"

    def _preprocess_summary_content(self, content: str, content_type: str) -> str:
        """预处理纪要内容，去除固定格式的标题"""
        if content_type == "editorjs":
            return content
        content = re.sub(r'<h2>\s*会议纪要\s*</h2>\s*', '', content)
        content = re.sub(r'## 会议纪要.*?### 关键词：\s*', '### 关键词：', content, flags=re.DOTALL)
        content = re.sub(r'<h2[^>]*>会议纪要.*?<h3[^>]*>关键词：\s*</h3>\s*', '<h3>关键词：</h3>', content, flags=re.DOTALL)
        content = re.sub(
            r'^会议纪要\s*会议基本信息\s*(•\s*会议时间：.*?\s*)*?(•\s*会议地点：.*?\s*)*?(•\s*主持人：.*?\s*)*?(•\s*参会人员：.*?\s*)*?',
            '', content, flags=re.DOTALL | re.MULTILINE)
        content = re.sub(
            r'^#\s*会议纪要\s*##\s*会议基本信息\s*(-\s*会议时间：.*?\s*)*?(-\s*会议地点：.*?\s*)*?(-\s*主持人：.*?\s*)*?(-\s*参会人员：.*?\s*)*?',
            '', content, flags=re.DOTALL | re.MULTILINE)
        return content

    def _build_summary_plain(self, content, meeting_data, content_type,
                             recording_time_str, attendees_str, output_path):
        """白板模式生成纪要"""
        theme = meeting_data.get('theme', '会议纪要')

        if content_type == "editorjs":
            editorjs_data = json.loads(content) if isinstance(content, str) else content
            meeting_info_blocks = [
                {"id": "m-theme", "type": "paragraph",
                 "data": {"text": f"<b>会议主题</b>：{meeting_data.get('theme', '无')}"}},
                {"id": "m-time", "type": "paragraph",
                 "data": {"text": f"<b>会议时间</b>：{recording_time_str or '无'}"}},
                {"id": "m-place", "type": "paragraph",
                 "data": {"text": f"<b>会议地点</b>：{meeting_data.get('places', '无')}"}},
                {"id": "m-moderator", "type": "paragraph",
                 "data": {"text": f"<b>主持人</b>：{meeting_data.get('moderator', '无')}"}},
                {"id": "m-attendees", "type": "paragraph",
                 "data": {"text": f"<b>参会人员</b>：{attendees_str}"}},
                {"id": "m-delimiter", "type": "delimiter", "data": {}},
            ]
            full_data = {
                "time": editorjs_data.get("time", int(datetime.now().timestamp() * 1000)),
                "blocks": meeting_info_blocks + editorjs_data.get("blocks", []),
                "version": editorjs_data.get("version", "2.31.3"),
            }
            return build_from_editor_js(full_data, output_path, f"{theme} - 会议纪要")

        meeting_info_html = f"""
        <p><strong>会议主题</strong>：{meeting_data.get('theme', '无')}</p>
        <p><strong>会议时间</strong>：{recording_time_str or '无'}</p>
        <p><strong>会议地点</strong>：{meeting_data.get('places', '无')}</p>
        <p><strong>主持人</strong>：{meeting_data.get('moderator', '无')}</p>
        <p><strong>参会人员</strong>：{attendees_str}</p>
        <hr/>
        """
        full_content = meeting_info_html + content
        doc_node = build_from_html(full_content, f"{theme} - 会议纪要")
        doc = Document()
        render_document(doc, doc_node)
        return save_docx(doc, output_path)

    def _build_summary_html_template(self, content, template_path, meeting_data,
                                     recording_time_str, attendees_str, output_path):
        """HTML/Markdown 模板模式生成纪要"""
        try:
            doc = load_template(template_path)
            placeholders = apply_mapping(meeting_data, MEETING_MAPPING)
            placeholders['会议时间'] = recording_time_str or '无'
            placeholders['参会人员'] = attendees_str
            replace_placeholders(doc, placeholders)

            doc_node = build_from_html(content)
            nodes = []
            for section in doc_node.sections:
                nodes.extend(section.body)
            replace_rich_placeholder(doc, '正文', nodes)

            return save_docx(doc, output_path)
        except Exception as e:
            logger.error(f"模板生成纪要失败: {e}")
            return False, None, str(e)

    def _build_summary_editorjs_template(self, content, template_path, meeting_data,
                                         recording_time_str, attendees_str, output_path):
        """Editor.js 模板模式生成纪要"""
        try:
            editorjs_data = json.loads(content) if isinstance(content, str) else content
            placeholders = {
                '会议主题': meeting_data.get('theme', '无'),
                '会议时间': recording_time_str or '无',
                '会议地点': meeting_data.get('places', '无'),
                '主持人': meeting_data.get('moderator', '无'),
                '参会人员': attendees_str,
                '会议类型': KIND_LABEL_MAP.get(meeting_data.get('kind', 0), '标准会议'),
                '备注': meeting_data.get('remark', '无'),
            }
            return convert_editorjs_with_template(
                editorjs_data, template_path,
                placeholders=placeholders,
                rich_placeholders={'正文': editorjs_data},
                output_path=output_path
            )
        except Exception as e:
            logger.error(f"Editor.js 模板生成纪要失败: {e}")
            return False, None, str(e)


# =============================================================================
# 便捷函数
# =============================================================================

def quick_build_note(content: str, output_path: str,
                     title: Optional[str] = None,
                     content_type: str = "auto") -> Tuple[bool, str, str]:
    """快速生成笔记文档"""
    gen = MeetingDocGenerator()
    return gen.build_note(content, title=title, content_type=content_type, output_path=output_path)


def quick_build_summary(content: str, output_path: str,
                        meeting_data: Optional[Dict[str, Any]] = None,
                        template_type: int = 1,
                        content_type: str = "auto") -> Tuple[bool, str, str]:
    """快速生成会议纪要"""
    gen = MeetingDocGenerator()
    return gen.build_summary(content, meeting_data=meeting_data,
                             template_type=template_type, content_type=content_type,
                             output_path=output_path)
