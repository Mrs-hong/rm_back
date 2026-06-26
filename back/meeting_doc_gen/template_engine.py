#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""会议文档生成器 - 模板引擎（占位符替换）"""

from typing import Dict, List, Optional

from docx import Document
from loguru import logger

from .models import ParagraphNode, TableNode
from .docx_renderer import apply_paragraph_format, render_run, render_table


def load_template(template_path: str) -> Document:
    """加载模板文件"""
    return Document(template_path)


def replace_placeholders(doc: Document, placeholders: Dict[str, str]):
    """替换文档中的普通占位符 {{key}} → value"""
    # 处理段落
    for para in doc.paragraphs:
        _replace_in_paragraph(para, placeholders)
    # 处理表格
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    _replace_in_paragraph(para, placeholders)


def _replace_in_paragraph(para, placeholders: Dict[str, str]):
    """替换单个段落中的占位符"""
    text = para.text
    if '{{' not in text:
        return
    new_text = text
    for key, value in placeholders.items():
        new_text = new_text.replace(f'{{{{{key}}}}}', str(value) if value else '')
    if new_text != text:
        if para.runs:
            for run in para.runs:
                run.text = ''
            para.runs[0].text = new_text
        else:
            para.text = new_text


def replace_rich_placeholder(doc: Document, placeholder: str, nodes: List):
    """替换富文本占位符 {{placeholder}} → 段落/表格节点列表"""
    if not nodes:
        return

    # 查找占位符段落
    target_para = None
    full_placeholder = f'{{{{{placeholder}}}}}'
    for para in doc.paragraphs:
        if full_placeholder in para.text:
            target_para = para
            break
    if target_para is None:
        return

    # 替换
    parent = target_para._element.getparent()
    para_index = list(parent).index(target_para._element)
    parent.remove(target_para._element)

    for node in reversed(nodes):
        if isinstance(node, TableNode):
            table_count_before = len(doc.tables)
            render_table(doc, node)
            if len(doc.tables) > table_count_before:
                new_table = doc.tables[-1]
                parent.insert(para_index, new_table._element)
        else:
            new_para = doc.add_paragraph()
            apply_paragraph_format(new_para, node)
            for run_node in node.runs:
                render_run(new_para, run_node)
            parent.insert(para_index, new_para._element)
