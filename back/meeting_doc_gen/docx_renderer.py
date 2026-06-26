#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""会议文档生成器 - DOCX 渲染器"""

import io
from typing import Any, List, Optional, Tuple

from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from loguru import logger

from .models import Alignment, DocumentNode, ParagraphNode, RunNode, SectionNode, TableNode


def render_document(doc: Document, doc_node: DocumentNode):
    """渲染文档结构树到 Word 文档"""
    for i, section_node in enumerate(doc_node.sections):
        if i > 0:
            doc.add_section()
        _render_section(doc, section_node)


def _render_section(doc: Document, section_node: SectionNode):
    """渲染版面"""
    section = doc.sections[-1] if doc.sections else doc.sections[0]
    try:
        section.page_width = Cm(section_node.page_width)
        section.page_height = Cm(section_node.page_height)
        section.top_margin = Cm(section_node.margin_top)
        section.bottom_margin = Cm(section_node.margin_bottom)
        section.left_margin = Cm(section_node.margin_left)
        section.right_margin = Cm(section_node.margin_right)
    except Exception as e:
        logger.warning(f"设置页面尺寸失败: {e}")

    for node in section_node.body:
        if isinstance(node, ParagraphNode):
            render_paragraph(doc, node)
        elif isinstance(node, TableNode):
            render_table(doc, node)


def render_paragraph(doc: Document, para_node: ParagraphNode):
    """渲染段落"""
    try:
        para = doc.add_paragraph()
        apply_paragraph_format(para, para_node)
        for run_node in para_node.runs:
            render_run(para, run_node)
    except Exception as e:
        logger.warning(f"渲染段落失败: {e}")


def apply_paragraph_format(para, para_node: ParagraphNode):
    """应用段落格式"""
    pf = para.paragraph_format
    if para_node.first_line_indent:
        pf.first_line_indent = Pt(para_node.first_line_indent)
    if para_node.left_indent:
        pf.left_indent = Pt(para_node.left_indent)
    if para_node.line_spacing:
        pf.line_spacing = para_node.line_spacing
    if para_node.space_before:
        pf.space_before = Pt(para_node.space_before)
    if para_node.space_after:
        pf.space_after = Pt(para_node.space_after)
    align_map = {
        Alignment.LEFT: WD_ALIGN_PARAGRAPH.LEFT,
        Alignment.CENTER: WD_ALIGN_PARAGRAPH.CENTER,
        Alignment.RIGHT: WD_ALIGN_PARAGRAPH.RIGHT,
        Alignment.JUSTIFY: WD_ALIGN_PARAGRAPH.JUSTIFY,
    }
    para.alignment = align_map.get(para_node.alignment, WD_ALIGN_PARAGRAPH.LEFT)


def render_run(para, run_node: RunNode):
    """渲染文本块"""
    try:
        run = para.add_run(run_node.text)
        run.font.name = run_node.font_name_ascii
        r_pr = run._element.rPr
        r_pr.rFonts.set(qn('w:eastAsia'), run_node.font_name)
        r_pr.rFonts.set(qn('w:ascii'), run_node.font_name_ascii)
        r_pr.rFonts.set(qn('w:hAnsi'), run_node.font_name_ascii)
        run.font.size = Pt(run_node.font_size)
        run.font.bold = run_node.bold
        run.font.italic = run_node.italic
        run.font.underline = run_node.underline
        if run_node.color and len(run_node.color) == 6:
            r, g, b = int(run_node.color[0:2], 16), int(run_node.color[2:4], 16), int(run_node.color[4:6], 16)
            run.font.color.rgb = RGBColor(r, g, b)
        if run_node.background_color and len(run_node.background_color) == 6:
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), run_node.background_color.upper())
            r_pr.append(shd)
    except Exception as e:
        logger.warning(f"渲染文本块失败: {e}")


def render_table(doc: Document, table_node: TableNode):
    """渲染表格"""
    try:
        if not table_node.rows:
            return
        row_count = len(table_node.rows)
        col_count = max(len(row) for row in table_node.rows) if table_node.rows else 0
        if col_count == 0:
            return
        table = doc.add_table(rows=row_count, cols=col_count)
        table.style = 'Table Grid'
        for i, row_data in enumerate(table_node.rows):
            for j, cell_data in enumerate(row_data):
                if j >= col_count:
                    continue
                cell = table.cell(i, j)
                if cell.paragraphs:
                    cell.paragraphs[0].clear()
                for para_idx, para_node in enumerate(cell_data.paragraphs):
                    if para_idx == 0 and cell.paragraphs:
                        para = cell.paragraphs[0]
                        apply_paragraph_format(para, para_node)
                    else:
                        para = cell.add_paragraph()
                        apply_paragraph_format(para, para_node)
                    for run_node in para_node.runs:
                        render_run(para, run_node)
    except Exception as e:
        logger.warning(f"渲染表格失败: {e}")


def save_docx(doc: Document, output_path: Optional[str]) -> Tuple[bool, Any, str]:
    """保存文档，返回 (success, result, error)"""
    try:
        if not output_path:
            buf = io.BytesIO()
            doc.save(buf)
            buf.seek(0)
            return True, buf.getvalue(), ''
        if not output_path.lower().endswith('.docx'):
            output_path = f"{output_path}.docx"
        doc.save(output_path)
        return True, output_path, ''
    except Exception as e:
        logger.error(f"保存文档失败: {e}")
        return False, None, str(e)
